"""
Document parser service for extracting raw text from uploaded files.

Supports PDF (via pdfplumber), DOCX (via python-docx), and plain TXT files.
Dispatches to the appropriate parser based on the file extension.
"""

from __future__ import annotations

import io
from pathlib import Path


class DocumentParsingError(Exception):
    """Raised when a document cannot be parsed due to corruption or an unreadable format."""


def parse_document(filename: str, content: bytes) -> str:
    """Dispatch to the appropriate parser based on the file extension.

    Args:
        filename: Original filename including extension.
        content:  Raw file bytes.

    Returns:
        Extracted text as a UTF-8 string.

    Raises:
        ValueError: If the file extension is not supported.
        DocumentParsingError: If the file is corrupted or unreadable.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(content)
    elif ext == ".docx":
        return _parse_docx(content)
    elif ext == ".txt":
        return _parse_txt(content)
    else:
        raise ValueError(f"Unsupported file type: {ext!r}")


def _parse_pdf(content: bytes) -> str:
    """Extract text from a PDF file page by page, joined with newlines.

    Args:
        content: Raw PDF bytes.

    Returns:
        Extracted text with pages separated by newline characters.

    Raises:
        DocumentParsingError: If the PDF cannot be opened or read.
    """
    try:
        import pdfplumber  # type: ignore
    except ImportError as exc:
        raise DocumentParsingError(
            "pdfplumber is not installed. Install it with: pip install pdfplumber"
        ) from exc

    try:
        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                pages.append(page_text)
        return "\n".join(pages)
    except DocumentParsingError:
        raise
    except Exception as exc:
        raise DocumentParsingError(
            f"Failed to parse PDF: {exc}"
        ) from exc


def _parse_docx(content: bytes) -> str:
    """Extract text from a DOCX file (paragraphs and table cells), joined with newlines.

    Args:
        content: Raw DOCX bytes.

    Returns:
        Extracted text with paragraphs and table cells separated by newline characters.

    Raises:
        DocumentParsingError: If the DOCX cannot be opened or read.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise DocumentParsingError(
            "python-docx is not installed. Install it with: pip install python-docx"
        ) from exc

    try:
        doc = Document(io.BytesIO(content))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)

        return "\n".join(parts)
    except DocumentParsingError:
        raise
    except Exception as exc:
        raise DocumentParsingError(
            f"Failed to parse DOCX: {exc}"
        ) from exc


def _parse_txt(content: bytes) -> str:
    """Decode plain-text file bytes as UTF-8.

    Args:
        content: Raw text file bytes.

    Returns:
        Decoded UTF-8 string.

    Raises:
        DocumentParsingError: If the bytes cannot be decoded as UTF-8.
    """
    try:
        return content.decode("utf-8")
    except (UnicodeDecodeError, ValueError) as exc:
        raise DocumentParsingError(
            f"Failed to decode TXT file as UTF-8: {exc}"
        ) from exc
